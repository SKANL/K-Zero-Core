"""Tools DOCX para lectura, creación y edición por copia."""
from __future__ import annotations

from typing import Any

from k_zero_core.core.tool_safety import resolve_safe_path
from k_zero_core.core.tools.document_common import (
    _created_result,
    _export_path,
    _first_heading,
    _hex_rgb,
    _table_rows,
    _theme_color,
)
from k_zero_core.services.design_md import aplicar_diseno_entregable


def _set_docx_cell_shading(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill.strip("#"))
    tc_pr.append(shading)


def _set_docx_font_color(run: Any, hex_color: str) -> None:
    from docx.shared import RGBColor

    red, green, blue = _hex_rgb(hex_color)
    run.font.color.rgb = RGBColor(red, green, blue)


def _configure_docx_styles(doc: Any, primary: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, bold in (("Title", 24, False), ("Subtitle", 11, False), ("Heading 1", 14, True), ("Heading 2", 12, True)):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(10 if "Heading" in style_name else 0)
        style.paragraph_format.space_after = Pt(5)
        if style_name in {"Title", "Subtitle"}:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_docx_cover(doc: Any, title: str, subtitle: str, primary: str) -> None:
    from datetime import date
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt

    title_paragraph = doc.add_paragraph(style="Title")
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(24)
    _set_docx_font_color(title_run, primary)

    if subtitle:
        subtitle_paragraph = doc.add_paragraph(style="Subtitle")
        subtitle_paragraph.add_run(subtitle)

    meta = doc.add_paragraph()
    meta.add_run("Entregable ejecutivo").bold = True
    meta.add_run(f"\nGenerado por K-Zero\n{date.today().isoformat()}")

    separator = doc.add_paragraph()
    separator_run = separator.add_run("-" * 58)
    _set_docx_font_color(separator_run, primary)
    separator.add_run().add_break(WD_BREAK.PAGE)


def _add_docx_sources(doc: Any, sources: list[str], primary: str) -> None:
    heading = doc.add_heading("Fuentes verificadas" if sources else "Fuentes no verificadas", level=1)
    for run in heading.runs:
        _set_docx_font_color(run, primary)
    if sources:
        for source in sources:
            doc.add_paragraph(source, style="List Bullet")
    else:
        doc.add_paragraph(
            "No se detectaron URLs verificables en el contenido entregado. "
            "Si se usó información externa, el productor debe volver a buscar y citar fuentes."
        )


def analizar_docx(path: str, max_chars: int = 12000) -> str:
    """Extrae texto y tablas básicas de un documento DOCX."""
    try:
        resolved = resolve_safe_path(path)
        from docx import Document

        doc = Document(str(resolved))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables[:5]:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows[:10]]
            tables.extend(rows)
        text = "\n".join(paragraphs + tables)
        return f"DOCX: {resolved}\nPárrafos: {len(paragraphs)}\nTablas: {len(doc.tables)}\n\n{text[:max_chars]}"
    except Exception as exc:
        return f"Error al analizar DOCX: {exc}"


def crear_docx(contenido_markdown: str, nombre_sugerido: str = "documento", design_md_path: str = "") -> str:
    """Crea un DOCX nuevo en exports con estilos nativos y diseño DESIGN.md."""
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import Inches

        deliverable = aplicar_diseno_entregable(contenido_markdown, "docx", design_md_path)
        blocks = deliverable["blocks"]
        primary = _theme_color(deliverable, "colors.primary", "#1f5f8b")
        header_text = _theme_color(deliverable, "colors.surface", "#ffffff")
        title = _first_heading(blocks, nombre_sugerido)
        subtitle = next((block["text"] for block in blocks if block["type"] == "paragraph"), "")
        doc = Document()
        doc.core_properties.title = title
        doc.core_properties.subject = "Entregable ejecutivo K-Zero"
        doc.core_properties.author = "K-Zero"
        doc.core_properties.keywords = "K-Zero, DESIGN.md, entregable, ejecutivo"
        section = doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        _configure_docx_styles(doc, primary)
        _add_docx_cover(doc, title, subtitle, primary)

        skipped_cover_title = False
        for block in blocks:
            kind = block["type"]
            if kind == "heading":
                level = int(block.get("level", 1))
                if level == 1 and block["text"] == title and not skipped_cover_title:
                    skipped_cover_title = True
                    continue
                paragraph = doc.add_heading(block["text"], level=0 if level == 1 else min(level - 1, 4))
                for run in paragraph.runs:
                    _set_docx_font_color(run, primary)
            elif kind == "paragraph":
                doc.add_paragraph(block["text"])
            elif kind == "bullets":
                for item in block["items"]:
                    doc.add_paragraph(item, style="List Bullet")
            elif kind == "numbers":
                for item in block["items"]:
                    doc.add_paragraph(item, style="List Number")
            elif kind == "table":
                rows = _table_rows(block)
                table = doc.add_table(rows=len(rows), cols=len(block["headers"]))
                table.style = "Table Grid"
                table.autofit = True
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for row_index, row in enumerate(rows):
                    for col_index, value in enumerate(row):
                        cell = table.cell(row_index, col_index)
                        cell.text = str(value)
                        if row_index == 0:
                            _set_docx_cell_shading(cell, primary)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    _set_docx_font_color(run, header_text)
                doc.add_paragraph("")
            elif kind == "code":
                doc.add_paragraph(block["text"], style="Intense Quote")
        _add_docx_sources(doc, deliverable["sources"], primary)
        output = _export_path(nombre_sugerido, ".docx")
        doc.save(str(output))
        return _created_result(output, "docx", nombre_sugerido)
    except Exception as exc:
        return f"Error al crear DOCX: {exc}"


def editar_docx_copia(path: str, instrucciones: str) -> str:
    """Crea una copia DOCX con una nota de edición; nunca modifica el original."""
    try:
        resolved = resolve_safe_path(path)
        from docx import Document

        doc = Document(str(resolved))
        doc.add_heading("Notas de edición solicitadas", level=1)
        doc.add_paragraph(instrucciones)
        output = _export_path(resolved.stem + "_editado", ".docx")
        doc.save(str(output))
        return f"Copia editada: {output}"
    except Exception as exc:
        return f"Error al editar DOCX: {exc}"
